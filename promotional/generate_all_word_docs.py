"""生成ThinkAi所有推广Word文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p

def set_paragraph_style(doc, text, bold=False, size=12, color=None, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def generate_blog_word():
    """生成掘金/CSDN博客版Word文档"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.3

    # 标题
    title = doc.add_heading('一行代码调用13个AI大模型，这个Python框架太香了！', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 引言
    doc.add_heading('引言', level=2)
    doc.add_paragraph('你是否曾经被AI大模型的接入折磨过？')
    items = [
        '配置OpenAI API，再配DeepSeek，又配通义千问...',
        '每个API格式都不一样，请求参数、响应解析全要重写',
        '想换个模型？改代码改到怀疑人生！'
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('今天给大家分享一个开源项目：ThinkAi，帮你终结这一切！')

    # 什么是ThinkAi
    doc.add_heading('什么是ThinkAi？', level=2)
    doc.add_paragraph('ThinkAi 是一个基于Python异步的企业级AI大模型集成框架，让你一行代码就能调用13个主流AI大模型。')

    add_code_block(doc, 'from thinkai import ThinkAI\n\nai = ThinkAI(provider="openai", api_key="your-key")\nresponse = await ai.chat("你好")\nprint(response.content)\n\n# 换成DeepSeek？只需改一行！\nai = ThinkAI(provider="deepseek", api_key="your-key")')
    doc.add_paragraph('就是这么简单！')

    # 核心优势
    doc.add_heading('核心优势', level=2)
    advantages = [
        ('一行代码接入', 'provider="模型名" 即可使用，无需重复配置'),
        ('13个主流模型', 'OpenAI、DeepSeek、通义千问、GLM、百度文心、腾讯混元、豆包、Kimi、MiniMax、Claude、Gemini、Ollama'),
        ('开箱即用', '内置所有API地址和默认模型，零配置'),
        ('轻量零依赖', '核心框架不依赖任何Web框架，完全独立'),
        ('高性能', '全异步架构 + 连接池 + 自动重试'),
        ('MIT开源免费', '完全免费，可商用'),
    ]
    for t, d in advantages:
        p = doc.add_paragraph()
        p.add_run(f'{t}：').bold = True
        p.add_run(d)

    # 支持模型
    doc.add_heading('支持哪些AI模型？', level=2)
    models = [
        ('Ollama', 'ollama', '本地运行，免费'),
        ('OpenAI GPT', 'openai', '国际主流'),
        ('DeepSeek', 'deepseek', '性价比高，免费额度'),
        ('通义千问', 'qwen', '阿里云出品'),
        ('智谱GLM', 'glm', '国产优秀'),
        ('百度文心', 'baidu', '中文理解强'),
        ('腾讯混元', 'tencent', '多模态支持'),
        ('豆包', 'doubao', '字节跳动'),
        ('Kimi', 'kimi', '长文本处理'),
        ('MiniMax', 'minimax', '多场景'),
        ('Claude', 'claude', 'Anthropic出品'),
        ('Gemini', 'gemini', 'Google出品'),
        ('自定义', 'openai-compatible', '任意OpenAI兼容API'),
    ]
    table = doc.add_table(rows=len(models)+1, cols=3, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['模型', 'Provider名称', '说明']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    for row_idx, (m, p, d) in enumerate(models, 1):
        table.rows[row_idx].cells[0].text = m
        table.rows[row_idx].cells[1].text = p
        table.rows[row_idx].cells[2].text = d

    # 快速开始
    doc.add_heading('快速开始', level=2)
    doc.add_heading('安装', level=3)
    add_code_block(doc, 'pip install thinkai-framework')
    
    doc.add_heading('3行代码开始使用', level=3)
    add_code_block(doc, 'from thinkai import ThinkAI\n\nai = ThinkAI(provider="ollama", model="llama3")\nresponse = await ai.chat("你好，世界")\nprint(response.content)')

    # 功能示例
    doc.add_heading('丰富功能示例', level=2)

    doc.add_heading('多模型自由切换', level=3)
    add_code_block(doc, 'from thinkai import ThinkAI\n\nai = ThinkAI(provider="openai", api_key="your-key")\nai.register_model("deepseek", provider="deepseek", api_key="your-key")\nai.register_model("qwen", provider="qwen", api_key="your-key")\n\nawait ai.chat("你好", model="openai")\nawait ai.chat("你好", model="deepseek")\nawait ai.chat("你好", model="qwen")')

    doc.add_heading('流式响应', level=3)
    add_code_block(doc, 'ai = ThinkAI(provider="openai", api_key="your-key")\n\nasync for chunk in ai.chat_stream("讲一个故事"):\n    if chunk.choices and chunk.choices[0].delta.content:\n        print(chunk.choices[0].delta.content, end="", flush=True)')

    doc.add_heading('会话管理（多轮对话）', level=3)
    add_code_block(doc, 'ai = ThinkAI()\n\nasync with ai.session() as session:\n    response1 = await session.chat("你好,我想学习Python")\n    response2 = await session.chat("有什么好的学习路径?")\n    response3 = await session.chat("推荐一些资源吧")')

    doc.add_heading('RAG（检索增强生成）', level=3)
    add_code_block(doc, 'from thinkai.rag import RAGPipeline\n\nrag = RAGPipeline(\n    documents=["./docs", "./knowledge"],\n    ai_client=ai,\n    chunk_size=500,\n)\n\nanswer = await rag.query("ThinkAi框架支持哪些AI模型?")')

    doc.add_heading('Agent（智能体）', level=3)
    add_code_block(doc, 'from thinkai.agent import ReActAgent, Tool\n\n@Tool(name="calculator", description="计算数学表达式")\ndef calculator(expression: str) -> str:\n    return str(eval(expression))\n\nagent = ReActAgent(tools=[calculator], ai_client=ai)\nresult = await agent.run("计算25*48的结果")')

    doc.add_heading('FastAPI集成', level=3)
    add_code_block(doc, 'from fastapi import FastAPI\nfrom thinkai import ThinkAI\n\napp = FastAPI()\nai = ThinkAI(provider="openai", api_key="your-key")\n\n@app.post("/chat")\nasync def chat(message: str):\n    response = await ai.chat(message)\n    return {"content": response.content}')

    # 为什么选择
    doc.add_heading('为什么选择ThinkAi？', level=2)
    reasons = [
        '开箱即用：无需配置API地址，内置所有主流模型的默认配置',
        '一行切换：provider="模型名" 即可在不同模型间切换',
        '轻量独立：不绑定任何Web框架，可搭配Flask、Django、FastAPI等使用',
        '企业级：异步架构、连接池、自动重试、中间件链',
        '完全免费：MIT开源协议，可商用'
    ]
    for i, reason in enumerate(reasons, 1):
        doc.add_paragraph(f'{i}. {reason}')

    # 相关链接
    doc.add_heading('相关链接', level=2)
    links = [
        ('PyPI', 'https://pypi.org/project/thinkai-framework/'),
        ('Gitee', 'https://gitee.com/hongxinge/think-ai'),
        ('版本', 'v0.2.1'),
        ('许可证', 'MIT'),
    ]
    for name, url in links:
        p = doc.add_paragraph()
        p.add_run(f'{name}: ').bold = True
        p.add_run(url)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('觉得好用的朋友别忘了给个Star！').bold = True

    output_path = 'G:/ThinkAi/promotional/博客-掘金CSDN.docx'
    doc.save(output_path)
    print(f'Word文档已保存: {output_path}')

def generate_zhihu_word():
    """生成知乎版Word文档"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.3

    # 标题
    title = doc.add_heading('一行代码调用13个AI大模型，这个Python框架太香了！', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph('还在手动配置每个AI模型的API？请求格式不同、响应解析不同、切换模型改代码改到崩溃？')
    doc.add_paragraph('今天我给大家分享一个开源项目：ThinkAi')
    doc.add_paragraph('一个基于Python异步的企业级AI大模型集成框架，让你一行代码就能调用13个主流AI大模型。')

    doc.add_heading('安装只需一行命令', level=2)
    add_code_block(doc, 'pip install thinkai-framework')

    doc.add_heading('使用只需3行代码', level=2)
    add_code_block(doc, 'from thinkai import ThinkAI\n\nai = ThinkAI(provider="openai", api_key="your-key")\nresponse = await ai.chat("你好")\nprint(response.content)')

    doc.add_heading('支持13个模型', level=2)
    models = [
        'Ollama - 本地运行，免费',
        'OpenAI GPT - 国际主流',
        'DeepSeek - 性价比高，免费额度',
        '通义千问 - 阿里云出品',
        '智谱GLM - 国产优秀',
        '百度文心 - 中文理解强',
        '腾讯混元 - 多模态支持',
        '豆包 - 字节跳动',
        'Kimi - 长文本处理',
        'MiniMax - 多场景',
        'Claude - Anthropic出品',
        'Gemini - Google出品',
        '任意OpenAI兼容API - 自定义',
    ]
    for m in models:
        doc.add_paragraph(m, style='List Bullet')

    doc.add_heading('为什么选择ThinkAi？', level=2)
    reasons = [
        ('开箱即用', '内置所有API地址和默认模型，零配置'),
        ('一行切换', 'provider="模型名" 即可在不同模型间切换'),
        ('轻量独立', '不绑定任何Web框架，可搭配Flask、Django、FastAPI等任何框架'),
        ('高性能', '全异步架构 + 连接池 + 自动重试'),
        ('完全免费', 'MIT开源协议，可商用'),
    ]
    for t, d in reasons:
        p = doc.add_paragraph()
        p.add_run(f'{t}：').bold = True
        p.add_run(d)

    doc.add_heading('代码示例', level=2)
    add_code_block(doc, '# 使用OpenAI\nai = ThinkAI(provider="openai", api_key="your-key")\n\n# 换成DeepSeek？只需改一行！\nai = ThinkAI(provider="deepseek", api_key="your-key")\n\n# 换成本地Ollama？还是改一行！\nai = ThinkAI(provider="ollama", model="llama3")')

    doc.add_heading('更多实用功能', level=2)

    doc.add_heading('多模型同时管理', level=3)
    add_code_block(doc, 'ai = ThinkAI(provider="openai", api_key="your-key")\nai.register_model("deepseek", provider="deepseek", api_key="your-key")\nai.register_model("qwen", provider="qwen", api_key="your-key")\n\nawait ai.chat("你好", model="openai")\nawait ai.chat("你好", model="deepseek")\nawait ai.chat("你好", model="qwen")')

    doc.add_heading('流式输出', level=3)
    add_code_block(doc, 'ai = ThinkAI(provider="openai", api_key="your-key")\n\nasync for chunk in ai.chat_stream("讲一个故事"):\n    if chunk.choices and chunk.choices[0].delta.content:\n        print(chunk.choices[0].delta.content, end="", flush=True)')

    doc.add_heading('多轮对话', level=3)
    add_code_block(doc, 'ai = ThinkAI()\n\nasync with ai.session() as session:\n    await session.chat("你好,我想学习Python")\n    await session.chat("有什么好的学习路径?")\n    await session.chat("推荐一些资源吧")')

    doc.add_heading('RAG检索增强生成', level=3)
    add_code_block(doc, 'from thinkai.rag import RAGPipeline\n\nrag = RAGPipeline(documents=["./docs"], ai_client=ai)\nanswer = await rag.query("你的问题")')

    doc.add_heading('FastAPI集成', level=3)
    add_code_block(doc, 'from fastapi import FastAPI\nfrom thinkai import ThinkAI\n\napp = FastAPI()\nai = ThinkAI(provider="openai", api_key="your-key")\n\n@app.post("/chat")\nasync def chat(message: str):\n    response = await ai.chat(message)\n    return {"content": response.content}')

    # 项目信息
    doc.add_heading('项目信息', level=2)
    links = [
        ('PyPI', 'https://pypi.org/project/thinkai-framework/'),
        ('Gitee', 'https://gitee.com/hongxinge/think-ai'),
        ('版本', 'v0.2.1'),
        ('许可证', 'MIT'),
    ]
    for name, url in links:
        p = doc.add_paragraph()
        p.add_run(f'{name}：').bold = True
        p.add_run(url)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('觉得好用的朋友别忘了给个Star！').bold = True

    output_path = 'G:/ThinkAi/promotional/知乎.docx'
    doc.save(output_path)
    print(f'Word文档已保存: {output_path}')

if __name__ == '__main__':
    generate_blog_word()
    generate_zhihu_word()
    print('所有Word文档生成完成！')
