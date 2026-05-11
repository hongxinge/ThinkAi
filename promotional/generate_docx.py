"""生成ThinkAi推广文案Word文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# ========== 页面设置 ==========
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# ========== 标题 ==========
title = doc.add_heading('ThinkAi 企业级AI框架推广文档', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x33, 0x66, 0xCC)

# 版本信息
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info.add_run(f'版本: v0.2.1  |  日期: {datetime.date.today().strftime("%Y-%m-%d")}  |  许可证: MIT')
info_run.font.size = Pt(10)
info_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_paragraph()  # 空行

# ========== 一、宣传标题 ==========
doc.add_heading('一、宣传标题', level=1)
doc.add_heading('主标题', level=2)
doc.add_paragraph('一行代码调用13个AI大模型！', style='Intense Quote')

doc.add_heading('副标题', level=2)
doc.add_paragraph('别再手动配置每个AI模型了！ThinkAi像用requests一样简单，开箱即用，MIT开源免费！', style='Intense Quote')

# ========== 二、完整推广文案 ==========
doc.add_heading('二、完整推广文案', level=1)

doc.add_heading('引言', level=2)
p = doc.add_paragraph('你是否曾经被AI大模型的接入折磨过？')
doc.add_paragraph('OpenAI配置一遍，再配DeepSeek，又配通义千问...', style='List Bullet')
doc.add_paragraph('每个API格式都不一样，请求参数、响应解析全要重写', style='List Bullet')
doc.add_paragraph('想换个模型？改代码改到怀疑人生！', style='List Bullet')
doc.add_paragraph('现在，ThinkAi 帮你终结这一切！', style='Intense Quote')

doc.add_heading('什么是ThinkAi？', level=2)
doc.add_paragraph('ThinkAi 是一个基于Python异步的企业级AI大模型集成框架，让你一行代码就能调用13个主流AI大模型。')

# 代码块
p = doc.add_paragraph()
p.style = doc.styles['Normal']
run = p.add_run('from thinkai import ThinkAI\n\nai = ThinkAI(provider="openai", api_key="your-key")\nresponse = await ai.chat("你好")\nprint(response.content)')
run.font.name = 'Consolas'
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)

doc.add_paragraph('就是这么简单！')

# ========== 核心优势 ==========
doc.add_heading('核心优势', level=2)
advantages = [
    ('一行代码接入', 'provider="模型名" 即可使用，无需重复配置'),
    ('13个主流模型', 'OpenAI、DeepSeek、通义千问、GLM、百度文心、腾讯混元、豆包、Kimi、MiniMax、Claude、Gemini、Ollama'),
    ('开箱即用', '内置所有API地址和默认模型，零配置'),
    ('轻量零依赖', '核心框架不依赖任何Web框架，完全独立'),
    ('高性能', '全异步架构 + 连接池 + 自动重试'),
    ('MIT开源免费', '完全免费，可商用'),
]

for title, desc in advantages:
    p = doc.add_paragraph()
    p.add_run(f'✅ {title}').bold = True
    p.add_run(f'：{desc}')

# ========== 支持的AI模型 ==========
doc.add_heading('支持的AI模型', level=2)
table = doc.add_table(rows=14, cols=3, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 表头
headers = ['模型', 'Provider名称', '说明']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)

# 表格数据
models_data = [
    ('Ollama', 'ollama', '本地运行，免费'),
    ('OpenAI GPT', 'openai', '国际主流'),
    ('DeepSeek', 'deepseek', '性价比高，免费额度'),
    ('通义千问', 'qwen', '阿里云出品'),
    ('智谱GLM', 'glm', '国产优秀'),
    ('百度文心', 'baidu', '中文理解强'),
    ('腾讯混元', 'tencent', '多模态支持'),
    ('豆包', 'doubao', '字节跳动'),
    ('Kimi', 'kimi', '长文本处理'),
    ('MiniMax', 'minimax', '多场景'),
    ('Claude', 'claude', 'Anthropic出品'),
    ('Gemini', 'gemini', 'Google出品'),
    ('自定义', 'openai-compatible', '任意OpenAI兼容API'),
]

for row_idx, (model, provider, desc) in enumerate(models_data, 1):
    table.rows[row_idx].cells[0].text = model
    table.rows[row_idx].cells[1].text = provider
    table.rows[row_idx].cells[2].text = desc
    for cell in table.rows[row_idx].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

# ========== 快速开始 ==========
doc.add_heading('快速开始', level=2)

p = doc.add_paragraph()
p.add_run('安装：').bold = True
p.add_run('\n')
run = p.add_run('pip install thinkai-framework')
run.font.name = 'Consolas'
run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

p = doc.add_paragraph()
p.add_run('使用：').bold = True
p.add_run('\n')
code_text = '''from thinkai import ThinkAI

ai = ThinkAI(provider="ollama", model="llama3")
response = await ai.chat("你好，世界")
print(response.content)'''
run = p.add_run(code_text)
run.font.name = 'Consolas'
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

# ========== 更多示例 ==========
doc.add_heading('更多示例', level=2)

doc.add_heading('多模型自由切换', level=3)
code_text = '''ai = ThinkAI(provider="openai", api_key="your-key")
ai.register_model("deepseek", provider="deepseek", api_key="your-key")
ai.register_model("qwen", provider="qwen", api_key="your-key")

await ai.chat("你好", model="openai")
await ai.chat("你好", model="deepseek")
await ai.chat("你好", model="qwen")'''
p = doc.add_paragraph()
run = p.add_run(code_text)
run.font.name = 'Consolas'
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

doc.add_heading('流式响应', level=3)
code_text = '''ai = ThinkAI(provider="openai", api_key="your-key")
async for chunk in ai.chat_stream("讲一个故事"):
    if chunk.choices and chunk.choices[0].delta.content:
        print(chunk.choices[0].delta.content, end="", flush=True)'''
p = doc.add_paragraph()
run = p.add_run(code_text)
run.font.name = 'Consolas'
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

# ========== 社交媒体短文案 ==========
doc.add_heading('三、社交媒体短文案', level=1)

doc.add_heading('微博/朋友圈版', level=2)
p = doc.add_paragraph('🤖 一行代码调用13个AI大模型！')
doc.add_paragraph('别再手动配置每个AI了！ThinkAi框架让你像用requests一样简单调用：')
models_list = '✅ OpenAI GPT\n✅ DeepSeek\n✅ 通义千问\n✅ GLM\n✅ 百度文心\n✅ 腾讯混元\n✅ 豆包\n✅ Kimi\n✅ MiniMax\n✅ Claude\n✅ Gemini\n✅ Ollama\n✅ 任意OpenAI兼容API'
doc.add_paragraph(models_list)
doc.add_paragraph('pip install thinkai-framework')
doc.add_paragraph('MIT开源，完全免费！')

doc.add_heading('掘金/知乎版', level=2)
p = doc.add_paragraph()
p.add_run('标题：').bold = True
p.add_run('一行代码调用13个AI大模型，这个Python框架太香了！')
doc.add_paragraph('还在手动配置每个AI模型的API？请求格式不同、响应解析不同、切换模型改代码改到崩溃？')
doc.add_paragraph('今天我给大家分享一个开源项目：ThinkAi')
doc.add_paragraph('一个基于Python异步的企业级AI大模型集成框架，让你一行代码就能调用13个主流AI大模型。')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('觉得好用的朋友别忘了给个Star！⭐').bold = True

# ========== 宣传海报文案 ==========
doc.add_heading('四、宣传海报文案', level=1)

doc.add_heading('海报主标题', level=2)
doc.add_paragraph('一行代码调用13个AI大模型', style='Intense Quote')

doc.add_heading('海报副标题', level=2)
doc.add_paragraph('ThinkAi - 像用requests一样简单', style='Intense Quote')

doc.add_heading('海报要点', level=2)
doc.add_paragraph('🎯 一行代码接入', style='List Bullet')
doc.add_paragraph('🌐 13个主流模型', style='List Bullet')
doc.add_paragraph('⚡ 开箱即用', style='List Bullet')
doc.add_paragraph('🆓 MIT开源免费', style='List Bullet')

doc.add_paragraph()

# ========== 相关链接 ==========
doc.add_heading('五、相关链接', level=1)

links = [
    ('PyPI', 'https://pypi.org/project/thinkai-framework/'),
    ('Gitee', 'https://gitee.com/hongxinge/think-ai'),
    ('版本', 'v0.2.1'),
    ('许可证', 'MIT'),
]

for name, url in links:
    p = doc.add_paragraph()
    p.add_run(f'{name}: ').bold = True
    p.add_run(url)

# 保存
output_path = 'G:/ThinkAi/promotional/promotion-content.docx'
doc.save(output_path)
print(f'Word文档已保存: {output_path}')
