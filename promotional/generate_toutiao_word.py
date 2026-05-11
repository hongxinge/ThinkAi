"""生成今日头条版Word文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# 标题
title = doc.add_heading('一行代码调用13个AI大模型！这个Python开源框架让开发者直呼内行', level=1)

# 引言
doc.add_heading('还在手动配置每个AI模型？你可能OUT了', level=2)

doc.add_paragraph('做AI应用开发的程序员朋友们，是不是经常遇到这样的烦恼：')

p = doc.add_paragraph()
p.add_run('今天要用OpenAI的GPT模型，去官网申请API Key，看文档，写代码配置。')

p = doc.add_paragraph()
p.add_run('明天老板说要换DeepSeek，因为性价比高，又要重新申请、重新配置。')

p = doc.add_paragraph()
p.add_run('后天客户说要用通义千问或者百度文心，因为数据要在国内...')

doc.add_paragraph('每个模型的API地址不一样，请求格式不一样，响应解析也不一样。换一次模型，代码就要改一遍。')

p = doc.add_paragraph()
p.add_run('如果有一个框架，能让你一行代码就切换不同的AI大模型，会不会很爽？').bold = True

doc.add_paragraph('今天给大家介绍的这个开源项目，就解决了这个问题。')

# ThinkAi是什么
doc.add_heading('ThinkAi是什么？', level=2)

doc.add_paragraph('ThinkAi是一个Python开源框架，专门用来简化AI大模型的调用。')
doc.add_paragraph('它的特点很简单：不管你用哪个AI模型，代码写法都一样。')

doc.add_paragraph('比如你要用OpenAI：')
add_code_block(doc, 'from thinkai import ThinkAI\n\nai = ThinkAI(provider="openai", api_key="你的密钥")\n回复 = await ai.chat("你好")\nprint(回复.content)')

doc.add_paragraph('如果第二天想换成DeepSeek，只需要改一行代码：')
add_code_block(doc, 'ai = ThinkAI(provider="deepseek", api_key="你的密钥")')

doc.add_paragraph('就这一行，其他代码完全不用动。')

# 支持哪些模型
doc.add_heading('支持哪些AI模型？', level=2)

doc.add_paragraph('目前支持13个主流AI大模型，基本覆盖了市面上所有常用的：')

models = [
    'OpenAI GPT - 国际最主流的AI模型',
    'DeepSeek - 国产黑马，免费额度多，性价比极高',
    '通义千问 - 阿里云出品，中文能力强',
    '智谱GLM - 国产优秀模型',
    '百度文心 - 中文理解能力强',
    '腾讯混元 - 多模态支持好',
    '豆包 - 字节跳动旗下',
    'Kimi - 长文本处理能力强',
    'MiniMax - 多场景适用',
    'Claude - Anthropic出品，代码能力强',
    'Gemini - Google出品',
    'Ollama - 本地部署，完全免费',
    '任意OpenAI兼容的API - 自定义接入',
]
for m in models:
    doc.add_paragraph(m, style='List Number')

# 解决了什么问题
doc.add_heading('这个框架解决了什么问题？', level=2)

doc.add_heading('1. 配置太麻烦', level=3)
doc.add_paragraph('以前每个模型都要自己找API地址、看文档、配置请求格式。现在ThinkAi内置了所有常用模型的配置，你只需要填API Key就行。')

doc.add_heading('2. 切换模型要改很多代码', level=3)
doc.add_paragraph('以前换模型要改API地址、请求参数、响应解析...现在只需要改一个参数：provider="模型名"')

doc.add_heading('3. 代码重复太多', level=3)
doc.add_paragraph('每个模型写一套调用代码，维护成本高。现在一套代码通吃所有模型。')

doc.add_heading('4. 没有统一的标准', level=3)
doc.add_paragraph('不同模型返回的数据格式不一样，处理起来麻烦。ThinkAi统一了所有模型的返回格式，你拿到的数据结构永远是一样的。')

# 使用场景
doc.add_heading('实际使用场景举例', level=2)

doc.add_heading('场景一：对比不同模型的回答效果', level=3)
add_code_block(doc, 'from thinkai import ThinkAI\n\nai = ThinkAI(provider="openai", api_key="密钥1")\nai.register_model("deepseek", provider="deepseek", api_key="密钥2")\nai.register_model("qwen", provider="qwen", api_key="密钥3")\n\n问题 = "帮我解释一下什么是递归"\n\n回答1 = await ai.chat(问题, model="openai")\n回答2 = await ai.chat(问题, model="deepseek")\n回答3 = await ai.chat(问题, model="qwen")')

doc.add_paragraph('三行代码，三个模型的回答都拿到了，方便对比。')

doc.add_heading('场景二：根据成本自动选择模型', level=3)
add_code_block(doc, '# 简单问题用便宜的模型\nif 问题.长度 < 20:\n    ai = ThinkAI(provider="deepseek", api_key="密钥")\n# 复杂问题用好模型\nelse:\n    ai = ThinkAI(provider="openai", api_key="密钥")')

doc.add_heading('场景三：本地免费模型先用', level=3)
add_code_block(doc, '# 本地先跑Ollama，不用花钱\nai = ThinkAI(provider="ollama", model="llama3")\n回复 = await ai.chat("你好")')

# 怎么安装
doc.add_heading('怎么安装和使用？', level=2)

doc.add_paragraph('安装非常简单，就一行命令：')
add_code_block(doc, 'pip install thinkai-framework')

doc.add_paragraph('然后就可以用了：')
add_code_block(doc, 'from thinkai import ThinkAI\n\n# 3行代码开始\nai = ThinkAI(provider="ollama", model="llama3")\n回复 = await ai.chat("你好，世界")\nprint(回复.content)')

# 其他功能
doc.add_heading('除了基础对话，还有什么功能？', level=2)

doc.add_heading('流式输出', level=3)
doc.add_paragraph('就是那种一个字一个字蹦出来的效果，用户体验更好。')
add_code_block(doc, 'ai = ThinkAI(provider="openai", api_key="密钥")\n\nasync for 片段 in ai.chat_stream("讲个故事"):\n    if 片段.choices and 片段.choices[0].delta.content:\n        print(片段.choices[0].delta.content, end="", flush=True)')

doc.add_heading('多轮对话', level=3)
doc.add_paragraph('就是让AI记住之前说过的话，进行连续对话。')
add_code_block(doc, 'ai = ThinkAI()\n\nasync with ai.session() as 会话:\n    await 会话.chat("你好，我想学习Python")\n    await 会话.chat("有什么好的学习路径？")\n    await 会话.chat("推荐一些资源吧")')

doc.add_heading('RAG（基于你的文档回答问题）', level=3)
doc.add_paragraph('把你的文档喂给AI，让它基于你的资料回答问题。')
add_code_block(doc, 'from thinkai.rag import RAGPipeline\n\nrag = RAGPipeline(文档路径=["./资料"], ai客户端=ai)\n答案 = await rag.query("你的问题")')

# 关于项目
doc.add_heading('关于这个项目', level=2)

info = [
    '完全开源：MIT协议，免费使用，可以商用',
    '代码仓库：https://gitee.com/hongxinge/think-ai',
    '安装包：https://pypi.org/project/thinkai-framework/',
    '当前版本：v0.2.1',
]
for item in info:
    doc.add_paragraph(item)

doc.add_paragraph('如果你觉得这个框架有用，可以去代码仓库点个Star支持一下。')

# 总结
doc.add_heading('总结一下', level=2)

doc.add_paragraph('ThinkAi解决了一个很实际的问题：让开发者不用为了用不同的AI大模型而反复改代码。')
doc.add_paragraph('一行代码切换模型，开箱即用，MIT开源免费。')
doc.add_paragraph('对于做AI应用开发的程序员来说，能省下不少时间和精力。')
doc.add_paragraph('如果你正在做AI相关的项目，不妨试试这个框架。')

output_path = 'G:/ThinkAi/promotional/今日头条.docx'
doc.save(output_path)
print(f'Word文档已保存: {output_path}')
